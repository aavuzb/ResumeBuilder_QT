"""
docx_engine.py
Builds a clean, professional resume .docx from a plain data dictionary,
and can also export that same resume straight to PDF.

Used by resume_builder.py (the Tkinter GUI) but has no GUI dependencies itself,
so it can also be scripted / reused on its own.
"""

import os
import shutil
import subprocess
import sys
import tempfile

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from i18n import DEFAULT_LANGUAGE, t
from resume_style import resolve_style

# ---- Palette (defaults — the actual per-resume palette is resolved from
# data["style"] via resume_style.resolve_style(); see build_resume()) ---------
ACCENT = RGBColor(0x17, 0x3A, 0x56)        # deep navy — headings / rules
ACCENT_SOFT = RGBColor(0x3E, 0x63, 0x82)   # secondary accent (company names)
DARK = RGBColor(0x21, 0x24, 0x28)          # near-black body text
GRAY = RGBColor(0x5B, 0x63, 0x6B)          # secondary / meta text
LINE = RGBColor(0xC9, 0xD2, 0xD8)          # hairline rules — not user-styled

PAGE_MARGIN_IN = 0.65
FONT_BODY = "Calibri"
FONT_HEAD = "Calibri"

# ---- Type scale (kept to five clean sizes for a consistent, disciplined look) --
SIZE_NAME = 20        # name at the top of the resume
SIZE_TITLE = 12        # professional title under the name
SIZE_HEADING = 11      # section headings (SUMMARY, EXPERIENCE, ...) and job titles
SIZE_BODY = 10         # bullets, company/school names, degrees, skills, summary
SIZE_META = 9          # dates, contact line, detail/meta sub-lines


def _hex_to_rgb(hex_color):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ---- Low-level helpers ------------------------------------------------------
def _set_border(paragraph, color="173A56", size=6, space=4, position="bottom"):
    """Add a border under (or above) a paragraph — used for headings and rules."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    edge = OxmlElement(f"w:{position}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    pBdr.append(edge)
    pPr.append(pBdr)


def _set_page_background(doc, hex_color):
    """Set the page/paper background fill (OOXML <w:background>) to match
    the live preview. Skipped for plain white — Word's default already is,
    and this avoids adding a "print background colors" surprise for the
    common case where nobody touched this setting.

    Renders correctly in the .docx when opened in Word, and in the live
    HTML preview. LibreOffice's headless --convert-to pdf, the PDF-export
    fallback on Linux, does not honor this element — a background other
    than white may come out white in that PDF even though the .docx itself
    is correct. The style dialog surfaces this caveat to the user."""
    hex_color = hex_color.lstrip("#")
    if hex_color.upper() == "FFFFFF":
        return
    background = OxmlElement("w:background")
    background.set(qn("w:color"), hex_color)
    doc.element.insert(0, background)


def _add_run(paragraph, text, size=10, bold=False, italic=False, color=DARK,
             spacing=None, font=FONT_BODY):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font
    if spacing is not None:
        rPr = run._r.get_or_add_rPr()
        spacing_el = OxmlElement("w:spacing")
        spacing_el.set(qn("w:val"), str(spacing))
        rPr.append(spacing_el)
    return run


def _add_hyperlink(paragraph, text, url, size=10, bold=False, italic=False, color=ACCENT_SOFT,
                    font=FONT_BODY):
    """Add a clickable hyperlink run to `paragraph`, styled to match the
    surrounding resume text (underlined, same accent color) rather than the
    default Word hyperlink blue."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)

    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    rPr.append(color_el)

    run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)
    return hyperlink


def _text_or_link(paragraph, text, url, size=10, bold=False, italic=False, color=DARK,
                   font=FONT_BODY):
    """Add `text` to `paragraph` as a hyperlink if `url` is set, otherwise as a plain run."""
    url = (url or "").strip()
    if url:
        _add_hyperlink(paragraph, text, url, size=size, bold=bold, italic=italic, color=color, font=font)
    else:
        _add_run(paragraph, text, size=size, bold=bold, italic=italic, color=color, font=font)


# ---- Section ordering ---------------------------------------------------
# Everyone writes a resume in a different order — some lead with Education,
# others with Publications. These are the sections between the fixed header
# (Personal Info) and the fixed footer (Save/Export), in the order they
# render by default; the UI lets the user rearrange them per-draft via
# data["section_order"].
DEFAULT_SECTION_ORDER = [
    "summary", "skills", "experience", "education",
    "publications", "certificates", "additional",
]


def resolve_section_order(data: dict):
    """The section order to actually render: the user's custom order
    (data["section_order"]), sanitized to known keys, with any missing
    section appended at the end in its default position — so a stale or
    partial order (e.g. from an older draft) never silently drops content."""
    requested = data.get("section_order") or []
    order = [k for k in requested if k in DEFAULT_SECTION_ORDER]
    order += [k for k in DEFAULT_SECTION_ORDER if k not in order]
    return order


# The four bits joined into the single contact line under the name/title.
# Reorderable via data["contact_order"], same sanitization approach as
# section order above.
DEFAULT_CONTACT_ORDER = ["location", "phone", "email", "visa"]


def resolve_contact_order(data: dict):
    requested = data.get("contact_order") or []
    order = [k for k in requested if k in DEFAULT_CONTACT_ORDER]
    order += [k for k in DEFAULT_CONTACT_ORDER if k not in order]
    return order


def format_date_range(start_date, end_date):
    """Shared by docx_engine and preview so an experience entry's Start
    Date / End Date fields always combine the same way, e.g. "Mar 2025 –
    Present". Falls back to whichever single date is present."""
    start = (start_date or "").strip()
    end = (end_date or "").strip()
    if start and end:
        return f"{start} – {end}"
    return start or end


def format_employment_type(key, lang=None):
    """Translate a stored employment-type key ("full_time") to display
    text; falls back to showing the raw value as-is for anything that
    isn't a recognized key (e.g. a hand-edited or very old draft)."""
    if not key:
        return ""
    lookup_key = f"employment.{key}"
    display = t(lookup_key, lang=lang)
    return display if display != lookup_key else key


# ---- Main builder ------------------------------------------------------------
def build_resume(data: dict, output_path: str):
    """
    data schema:
    {
      "personal": {"name", "title", "location", "phone", "email", "visa"},
      "links": [{"label", "value"}, ...]  (e.g. LinkedIn, GitHub, Portfolio — shown
                as clickable links under the contact line),
      "section_order": [str, ...]  (optional — subset/permutation of
                DEFAULT_SECTION_ORDER; see resolve_section_order()),
      "contact_order": [str, ...]  (optional — subset/permutation of
                DEFAULT_CONTACT_ORDER; see resolve_contact_order()),
      "summary": str,
      "skills": [{"label", "value"}, ...],
      "experience": [
          {"title", "company", "start_date", "end_date", "employment_type"
           (one of "", "Full-time", "Part-time", "Contract", "Freelance"),
           "projects": [{"name", "bullets": [str, ...]}, ...]}
      ],
      "education": [{"degree", "school", "meta"}, ...],
      "publications": [{"title", "detail"}, ...]  (also accepts plain strings),
      "certificates": [{"label", "value"}, ...],
      "additional": [{"label", "value"}, ...],
      "style": {see resume_style.DEFAULT_STYLE}  (optional \u2014 font, sizes,
                colors, page background; see resume_style.resolve_style()),
    }
    """
    doc = Document()

    lang = data.get("language") or DEFAULT_LANGUAGE
    style = resolve_style(data)
    font = style["font_family"]
    size_name = style["size_name"]
    size_title = style["size_title"]
    size_heading = style["size_heading"]
    size_body = style["size_body"]
    size_meta = style["size_meta"]
    accent = _hex_to_rgb(style["accent_color"])
    accent_soft = _hex_to_rgb(style["accent_soft_color"])
    highlight_hex = style["highlight_color"].lstrip("#")
    text_color = _hex_to_rgb(style["text_color"])
    muted = _hex_to_rgb(style["muted_color"])

    # Page setup
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(PAGE_MARGIN_IN)
    section.bottom_margin = Inches(PAGE_MARGIN_IN)
    section.left_margin = Inches(PAGE_MARGIN_IN)
    section.right_margin = Inches(PAGE_MARGIN_IN)
    usable_width = 8.5 - 2 * PAGE_MARGIN_IN

    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(size_body)
    try:
        bullet_style = doc.styles["List Bullet"]
        bullet_style.font.size = Pt(size_body)
        bullet_style.font.name = font
    except KeyError:
        pass
    _set_page_background(doc, style["background_color"])

    # ---- Section-body helpers \u2014 closures bound to this resume's style, so
    # every call below automatically picks up the chosen font/sizes/colors ----
    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        _add_run(p, text.upper(), size=size_heading, bold=True, color=accent, spacing=24, font=font)
        _set_border(p, color=highlight_hex, size=6, space=5)
        return p

    def bullet(text, indent=0.20):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.left_indent = Inches(indent)
        _add_run(p, text, size=size_body, color=text_color, font=font)
        return p

    def project_heading(name, url=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.10)
        _add_run(p, "\u2023  ", size=size_body, bold=True, color=accent_soft, font=font)
        _text_or_link(p, name, url, size=size_body, bold=True, color=text_color, font=font)
        return p

    def label_value(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.08
        _add_run(p, label + ":  ", size=size_body, bold=True, color=accent_soft, font=font)
        _add_run(p, value, size=size_body, color=text_color, font=font)
        return p

    def job_header(title, company, dates, company_url="", employment_type=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(usable_width), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
        )
        _add_run(p, title, size=size_heading, bold=True, color=accent, font=font)
        if employment_type:
            _add_run(p, f"   ({employment_type})", size=size_meta, italic=True, color=muted, font=font)
        _add_run(p, "\t" + dates, size=size_meta, italic=True, color=muted, font=font)
        if company:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(3)
            _text_or_link(p2, company, company_url, size=size_body, bold=True, italic=True,
                          color=accent_soft, font=font)
        return p

    def publication_entry(item):
        if isinstance(item, dict):
            title = (item.get("title") or "").strip()
            detail = (item.get("detail") or "").strip()
        else:
            title = str(item).strip()
            detail = ""
        if not title and not detail:
            return
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1 if detail else 4)
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.left_indent = Inches(0.20)
        _add_run(p, title, size=size_body, bold=True, color=text_color, font=font)
        if detail:
            dp = doc.add_paragraph()
            dp.paragraph_format.left_indent = Inches(0.20)
            dp.paragraph_format.space_after = Pt(6)
            _add_run(dp, detail, size=size_meta, italic=True, color=muted, font=font)

    personal = data.get("personal", {})

    # ---- Header ----
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    _add_run(name_p, personal.get("name", "").upper(), size=size_name, bold=True, color=accent,
             spacing=14, font=font)

    if personal.get("title"):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(7)
        _add_run(title_p, personal["title"], size=size_title, italic=True, color=accent_soft, font=font)

    contact_bits = [personal.get(k) for k in resolve_contact_order(data) if personal.get(k)]
    link_rows = [r for r in data.get("links", []) if r.get("label") or r.get("value")]

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(2 if link_rows else 12)
    _add_run(contact_p, "   \u2022   ".join(contact_bits), size=size_meta, color=muted, font=font)
    if not link_rows:
        _set_border(contact_p, color="C9D2D8", size=6, space=8)

    if link_rows:
        links_p = doc.add_paragraph()
        links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_p.paragraph_format.space_after = Pt(12)
        for i, row in enumerate(link_rows):
            if i > 0:
                _add_run(links_p, "   \u2022   ", size=size_meta, color=muted, font=font)
            _text_or_link(links_p, row.get("label", ""), row.get("value", ""),
                          size=size_meta, color=accent_soft, font=font)
        _set_border(links_p, color="C9D2D8", size=6, space=8)

    # ---- Section bodies (dispatched below in the user's chosen order) ----
    def render_summary():
        if not data.get("summary"):
            return
        heading(t("resume.summary", lang=lang))
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        _add_run(p, data["summary"], size=size_body, color=text_color, font=font)

    def render_skills():
        if not data.get("skills"):
            return
        heading(t("resume.skills", lang=lang))
        for row in data["skills"]:
            if row.get("label") or row.get("value"):
                label_value(row.get("label", ""), row.get("value", ""))

    def render_experience():
        if not data.get("experience"):
            return
        heading(t("resume.experience", lang=lang))
        for job in data["experience"]:
            dates_display = format_date_range(job.get("start_date", ""), job.get("end_date", ""))
            if not dates_display and job.get("dates"):
                dates_display = job["dates"]  # older drafts stored one combined string
            employment_display = format_employment_type(job.get("employment_type", ""), lang)
            job_header(job.get("title", ""), job.get("company", ""), dates_display,
                       job.get("company_url", ""), employment_display)
            projects = job.get("projects")
            if projects:
                for proj in projects:
                    name = (proj.get("name") or "").strip()
                    if name:
                        project_heading(name, proj.get("url", ""))
                    for b in proj.get("bullets", []):
                        if b.strip():
                            bullet(b.strip(), indent=0.32 if name else 0.20)
            else:
                for b in job.get("bullets", []):
                    if b.strip():
                        bullet(b.strip())

    def render_education():
        if not data.get("education"):
            return
        heading(t("resume.education", lang=lang))
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _add_run(p, edu.get("degree", ""), size=size_body, bold=True, color=text_color, font=font)
            if edu.get("school"):
                _add_run(p, "  \u2014  ", size=size_body, color=muted, font=font)
                _text_or_link(p, edu["school"], edu.get("school_url", ""), size=size_body,
                              color=muted, font=font)
            if edu.get("meta"):
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.space_after = Pt(8)
                _add_run(meta_p, edu["meta"], size=size_meta, italic=True, color=muted, font=font)

    def render_publications():
        if not data.get("publications"):
            return
        heading(t("resume.publications", lang=lang))
        for item in data["publications"]:
            publication_entry(item)

    def render_certificates():
        if not data.get("certificates"):
            return
        heading(t("resume.certificates", lang=lang))
        for row in data["certificates"]:
            if row.get("label") or row.get("value"):
                label_value(row.get("label", ""), row.get("value", ""))

    def render_additional():
        if not data.get("additional"):
            return
        heading(t("resume.additional", lang=lang))
        for row in data["additional"]:
            if row.get("label") or row.get("value"):
                label_value(row.get("label", ""), row.get("value", ""))

    section_renderers = {
        "summary": render_summary,
        "skills": render_skills,
        "experience": render_experience,
        "education": render_education,
        "publications": render_publications,
        "certificates": render_certificates,
        "additional": render_additional,
    }
    for key in resolve_section_order(data):
        section_renderers[key]()

    doc.save(output_path)
    return output_path


# ---- PDF export ---------------------------------------------------------
class PDFConversionError(RuntimeError):
    """Raised when the resume could not be converted to PDF on this machine."""


def _convert_with_word_com(docx_path, pdf_path):
    """Windows only — uses installed Microsoft Word via COM automation."""
    import comtypes.client  # type: ignore

    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
    finally:
        word.Quit()


def _convert_with_docx2pdf(docx_path, pdf_path):
    """Windows or macOS with Word/AppleScript available."""
    from docx2pdf import convert  # type: ignore
    convert(docx_path, pdf_path)


def _convert_with_libreoffice(docx_path, pdf_path):
    """Cross-platform fallback — requires a LibreOffice/soffice install."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise PDFConversionError("LibreOffice (soffice) was not found on this system.")

    out_dir = tempfile.mkdtemp(prefix="resume_pdf_")
    result = subprocess.run(
        [soffice, "--headless", "--norestore", "--convert-to", "pdf",
         "--outdir", out_dir, os.path.abspath(docx_path)],
        capture_output=True, text=True, timeout=90,
    )
    produced = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if result.returncode != 0 or not os.path.exists(produced):
        raise PDFConversionError(
            "LibreOffice failed to convert the document to PDF.\n" + (result.stderr or "")
        )
    shutil.move(produced, pdf_path)
    shutil.rmtree(out_dir, ignore_errors=True)


def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """
    Convert an existing .docx file to .pdf, trying the best available method
    for the current platform, in order:
      1. Microsoft Word (Windows COM automation) — highest fidelity
      2. docx2pdf (Windows/macOS, drives Word or uses macOS's built-in engine)
      3. LibreOffice / soffice headless conversion (Windows/macOS/Linux)
    Raises PDFConversionError with a clear message if none succeed.
    """
    errors = []

    if sys.platform.startswith("win"):
        try:
            _convert_with_word_com(docx_path, pdf_path)
            return pdf_path
        except Exception as exc:
            errors.append(f"Word (COM): {exc}")

    if sys.platform.startswith("win") or sys.platform == "darwin":
        try:
            _convert_with_docx2pdf(docx_path, pdf_path)
            return pdf_path
        except Exception as exc:
            errors.append(f"docx2pdf: {exc}")

    try:
        _convert_with_libreoffice(docx_path, pdf_path)
        return pdf_path
    except Exception as exc:
        errors.append(f"LibreOffice: {exc}")

    raise PDFConversionError(
        "Could not convert to PDF using any available method.\n"
        "Install Microsoft Word (Windows/macOS) or LibreOffice "
        "(https://www.libreoffice.org/download) and try again.\n\n"
        + "\n".join(errors)
    )


def build_resume_pdf(data: dict, pdf_output_path: str):
    """Build the resume directly to PDF (via a temporary .docx)."""
    with tempfile.TemporaryDirectory() as tmp:
        tmp_docx = os.path.join(tmp, "resume_temp.docx")
        build_resume(data, tmp_docx)
        convert_docx_to_pdf(tmp_docx, pdf_output_path)
    return pdf_output_path
